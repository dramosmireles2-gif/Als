#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generador del Manual de Código — Als Dress
Produce: /home/user/Als/manual-codigo.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_font(run, name, size=None, bold=False, italic=False, color=None):
    run.font.name = name
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def h1(doc, text):
    p = doc.add_paragraph()
    p.style = 'Heading 1'
    run = p.runs[0] if p.runs else p.add_run(text)
    if not p.runs:
        run = p.add_run(text)
    else:
        run.text = text
    run.font.name = 'Calibri'
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1c, 0x1c, 0x1e)
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    p.style = 'Heading 2'
    run = p.runs[0] if p.runs else p.add_run(text)
    if not p.runs:
        run = p.add_run(text)
    else:
        run.text = text
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xc0, 0x27, 0x5a)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p

def h3(doc, text):
    p = doc.add_paragraph()
    p.style = 'Heading 3'
    run = p.runs[0] if p.runs else p.add_run(text)
    if not p.runs:
        run = p.add_run(text)
    else:
        run.text = text
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1c, 0x1c, 0x1e)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p

def body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(16)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(3)
    return p

def code_block(doc, code_text):
    """Bloque de código con fondo gris claro y fuente Courier."""
    # Paragraph shade
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    # indent
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return p

def separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('─' * 70)
    run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
    run.font.size = Pt(9)
    return p

def note_box(doc, text):
    """Cuadro de nota con borde izquierdo simulado."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF3F6')
    pPr.append(shd)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('📌 NOTA: ' + text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0xc0, 0x27, 0x5a)
    return p

# ── Documento ─────────────────────────────────────────────────────────────────

doc = Document()

# Márgenes
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.1)
section.right_margin  = Inches(1.1)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ══════════════════════════════════════════════════════════════════════════════
#  PORTADA
# ══════════════════════════════════════════════════════════════════════════════
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_title.add_run('ALS DRESS')
r.font.name = 'Calibri'
r.font.size = Pt(36)
r.font.bold = True
r.font.color.rgb = RGBColor(0xc0, 0x27, 0x5a)
p_title.paragraph_format.space_before = Pt(60)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_sub.add_run('Manual de Código Fuente')
r2.font.name = 'Calibri'
r2.font.size = Pt(24)
r2.font.color.rgb = RGBColor(0x1c, 0x1c, 0x1e)

p_sub2 = doc.add_paragraph()
p_sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p_sub2.add_run('Guía completa para entender, mantener y modificar\nel sitio web de renta de vestidos')
r3.font.name = 'Calibri'
r3.font.size = Pt(13)
r3.font.italic = True
r3.font.color.rgb = RGBColor(0x8a, 0x8a, 0x8e)
p_sub2.paragraph_format.space_before = Pt(10)

p_meta = doc.add_paragraph()
p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p_meta.add_run('Tecnologías: HTML5 · CSS3 · JavaScript ES2022 · Supabase · Netlify\nReynosa, Tamaulipas · 2026')
r4.font.name = 'Calibri'
r4.font.size = Pt(11)
r4.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)
p_meta.paragraph_format.space_before = Pt(30)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  ÍNDICE (manual)
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, 'Tabla de Contenido')

toc_items = [
    ('1.', 'Resumen General del Proyecto', '3'),
    ('2.', 'Estructura HTML — Secciones de la página', '7'),
    ('   2.1', 'Header (Cabecera)', '7'),
    ('   2.2', 'Hero (Portada principal)', '8'),
    ('   2.3', 'Categorías Visuales', '9'),
    ('   2.4', 'Catálogo de productos', '10'),
    ('   2.5', 'Proceso "Cómo funciona"', '11'),
    ('   2.6', 'Galería de Clientas', '12'),
    ('   2.7', 'Testimonios', '13'),
    ('   2.8', 'Ubicación', '14'),
    ('   2.9', 'Footer y botones flotantes', '15'),
    ('3.', 'Estilos CSS — Variables, tipografías y componentes', '16'),
    ('   3.1', 'Variables de color y tokens de diseño', '16'),
    ('   3.2', 'Tipografías', '17'),
    ('   3.3', 'Hero layout', '18'),
    ('   3.4', 'Cards del catálogo', '19'),
    ('   3.5', 'Header fijo y menú responsive', '20'),
    ('   3.6', 'Botones', '21'),
    ('   3.7', 'Animaciones y transiciones', '22'),
    ('   3.8', 'Responsive / Media queries', '23'),
    ('4.', 'Lógica JavaScript — Funciones y eventos', '24'),
    ('   4.1', 'Variables globales y constantes', '24'),
    ('   4.2', 'cargarInventario()', '25'),
    ('   4.3', 'renderizarCatalogo()', '26'),
    ('   4.4', 'crearLightbox() y abrirLightbox()', '28'),
    ('   4.5', 'cargarGaleriaClientas()', '30'),
    ('   4.6', 'filtrarPorCategoria y filtros de talla', '31'),
    ('   4.7', 'Event listeners del menú y scroll', '32'),
    ('   4.8', 'Reveal on scroll — IntersectionObserver', '33'),
    ('   4.9', 'Galería de interiores (slider)', '34'),
    ('   4.10', 'DOMContentLoaded — Inicialización', '35'),
    ('5.', 'Guía de cambios comunes', '36'),
    ('6.', 'Glosario de términos técnicos', '43'),
]

for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r_num = p.add_run(f'{num}  ')
    r_num.font.name = 'Calibri'
    r_num.font.size = Pt(10)
    r_num.font.bold = True
    r_num.font.color.rgb = RGBColor(0xc0, 0x27, 0x5a)
    r_tit = p.add_run(title)
    r_tit.font.name = 'Calibri'
    r_tit.font.size = Pt(10)
    r_tit.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CAPÍTULO 1 — RESUMEN GENERAL
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, 'Capítulo 1: Resumen General del Proyecto')

body(doc,
    'Als Dress es una página web para el negocio de renta y venta de vestidos de fiesta ubicado '
    'en Reynosa, Tamaulipas. El sitio permite a las clientas explorar el catálogo de vestidos, '
    'accesorios, zapatos y bolsas, consultar disponibilidad, conocer el proceso de renta y '
    'ponerse en contacto por WhatsApp directamente desde el navegador.')

body(doc,
    'El proyecto fue construido con tecnologías web estándar: HTML5 para la estructura, '
    'CSS3 para los estilos visuales y JavaScript moderno (ES2022) para la interactividad. '
    'Los datos del inventario se almacenan en Supabase, una base de datos en la nube, '
    'y el sitio se despliega automáticamente en Netlify.')

h2(doc, '1.1 Tecnologías utilizadas')

body(doc, 'El proyecto usa únicamente tres archivos principales que tú debes conocer:')
bullet(doc, 'index.html — La estructura de la página (el "esqueleto").')
bullet(doc, 'style.css — Los estilos visuales (colores, tamaños, posiciones).')
bullet(doc, 'main.js — La lógica interactiva (cargar datos, filtros, lightbox).')
bullet(doc, 'supabase.js — Archivo de configuración de la conexión a la base de datos (no se documenta aquí en detalle).')

body(doc, 'Herramientas externas (CDN y servicios):')
bullet(doc, 'Supabase — Base de datos PostgreSQL en la nube. Guarda el inventario de vestidos y la galería de clientas.')
bullet(doc, 'Netlify — Servicio de hosting. Publica el sitio automáticamente cada vez que haces un cambio en el repositorio.')
bullet(doc, 'Google Fonts — Fuentes tipográficas: Playfair Display (títulos elegantes) y Lato (texto del cuerpo).')
bullet(doc, 'WhatsApp API — Enlace especial (wa.me) que abre WhatsApp con un mensaje predefinido.')
bullet(doc, 'Google Maps Embed — Mapa interactivo incrustado en la sección de ubicación.')

h2(doc, '1.2 ¿Cómo está organizada la página?')

body(doc,
    'La página es un "single page" o página de una sola hoja. Esto significa que todo el contenido '
    'está en un solo archivo HTML y el usuario navega desplazándose hacia abajo. Los enlaces del '
    'menú (como "#catalogo" o "#proceso") no abren páginas nuevas, sino que desplazan la vista '
    'a esa sección de la misma página.')

body(doc, 'El orden de las secciones de arriba hacia abajo es:')
bullet(doc, '1. Header (menú de navegación fijo en la parte superior)')
bullet(doc, '2. Hero (portada grande con la imagen principal y los botones de llamado a la acción)')
bullet(doc, '3. Categorías Visuales (tarjetas con fotos para ir a Vestidos, Zapatos, Bolsas, Accesorios)')
bullet(doc, '4. Catálogo (grid de productos cargados desde Supabase con filtros)')
bullet(doc, '5. Proceso "Cómo Funciona" (los 3 pasos de la renta)')
bullet(doc, '6. Galería de Clientas (cinta infinita con fotos de clientas reales, visible solo si hay datos)')
bullet(doc, '7. Testimonios (opiniones de clientas verificadas)')
bullet(doc, '8. Ubicación (galería del local + mapa de Google Maps)')
bullet(doc, '9. Footer (links a redes sociales y créditos)')
bullet(doc, '10. Botones flotantes (WhatsApp y Facebook siempre visibles en la esquina)')

h2(doc, '1.3 Flujo general de carga de la página')

body(doc,
    'Cuando una clienta abre el sitio en su navegador, sucede lo siguiente en orden:')
bullet(doc, '1. El navegador descarga index.html y empieza a mostrar el contenido.')
bullet(doc, '2. Se carga style.css y se aplican todos los estilos visuales.')
bullet(doc, '3. Se carga la librería de Supabase desde una CDN (red de distribución).')
bullet(doc, '4. Se carga supabase.js que conecta con la base de datos.')
bullet(doc, '5. Se carga main.js que ejecuta toda la lógica interactiva.')
bullet(doc, '6. DOMContentLoaded se dispara: se inicializa el lightbox, se carga el inventario y la galería de clientas.')
bullet(doc, '7. cargarInventario() hace una consulta a Supabase y cuando recibe los datos, llama a renderizarCatalogo().')
bullet(doc, '8. Los productos aparecen uno a uno con una pequeña animación de entrada.')

h2(doc, '1.4 Estructura de archivos del proyecto')

code_block(doc,
'/home/user/Als/\n'
'├── index.html          ← Página principal\n'
'├── style.css           ← Todos los estilos\n'
'├── main.js             ← Toda la lógica JS\n'
'├── supabase.js         ← Configuración de Supabase\n'
'├── hero2.png           ← Imagen de la portada\n'
'├── Icono.jpg           ← Logo del negocio\n'
'├── cat-vestidos.png    ← Imagen de la tarjeta de categoría Vestidos\n'
'├── cat-zapatos.png     ← Imagen de la tarjeta de categoría Zapatos\n'
'├── cat-bolsas.png      ← Imagen de la tarjeta de categoría Bolsas\n'
'├── cat-accesorios.png  ← Imagen de la tarjeta de categoría Accesorios\n'
'├── interior1.jpeg      ← Foto del local (galería en Ubicación)\n'
'├── interior2.jpeg      ← Foto del local\n'
'└── interior3.jpeg      ← Foto del local')

h2(doc, '1.5 Tabla de datos en Supabase')

body(doc,
    'La base de datos en Supabase tiene al menos dos tablas importantes:')

body(doc, 'Tabla "inventario":')
bullet(doc, 'nombre — Nombre del vestido o producto (ej: "Vestido Rojo Coral")')
bullet(doc, 'tipo — Categoría: "Vestido", "Zapato", "Bolsa" o "Accesorios"')
bullet(doc, 'talla — Talla del artículo (ej: "S", "M", "XL")')
bullet(doc, 'estado_actual — "Disponible" o cualquier otro valor = "Rentado"')
bullet(doc, 'foto — Nombre del archivo de foto o URL completa')
bullet(doc, 'publicado — true/false: solo se muestran los artículos con publicado=true')

body(doc, 'Tabla "galeria_clientas":')
bullet(doc, 'foto_url — URL directa de la imagen de la clienta')
bullet(doc, 'nombre — Nombre de la clienta (opcional)')
bullet(doc, 'activa — true/false: solo se muestran las fotos activas')
bullet(doc, 'created_at — Fecha de carga (se ordenan de más nueva a más antigua)')

note_box(doc,
    'Para agregar o modificar productos en el catálogo, no es necesario tocar el código. '
    'Simplemente entra al dashboard de Supabase (supabase.com), abre el proyecto, ve a la '
    'tabla "inventario" y edita los registros ahí.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CAPÍTULO 2 — ESTRUCTURA HTML
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, 'Capítulo 2: Estructura HTML — Secciones de la Página')

body(doc,
    'HTML (HyperText Markup Language) es el lenguaje que define la estructura de la página. '
    'Funciona con "etiquetas" que envuelven el contenido. Por ejemplo, <h1> es un título principal, '
    '<p> es un párrafo, <section> es una sección de contenido, y <div> es un contenedor genérico. '
    'A continuación se explica cada sección del archivo index.html.')

h2(doc, '2.1 El documento HTML — <head>')

body(doc,
    'El <head> es la parte del HTML que el usuario NO ve directamente, pero que es fundamental. '
    'Contiene información sobre la página (metadatos) y los enlaces a los recursos externos.')

code_block(doc,
'<head>\n'
'  <meta charset="UTF-8">\n'
'  <meta name="viewport" content="width=device-width, initial-scale=1.0">\n'
'  <meta name="description" content="Als Dress — Renta y venta de vestidos...">\n'
'  <title>Als Dress | Renta de Vestidos en Reynosa</title>\n'
'  ...\n'
'  <link href="https://fonts.googleapis.com/css2?family=Playfair+Display..." rel="stylesheet">\n'
'  <link rel="stylesheet" href="style.css">\n'
'  <script src="https://cdn.jsdelivr.net/npm/@supabase/supabase-js@2"></script>\n'
'  <script src="./supabase.js"></script>\n'
'</head>')

body(doc, 'Explicación de cada parte:')
bullet(doc, 'charset="UTF-8" — Permite usar caracteres especiales como acentos (á, é, í) y la ñ.')
bullet(doc, 'viewport — Hace que la página se adapte correctamente a pantallas de celular.')
bullet(doc, 'description — Texto que aparece en Google cuando alguien busca el negocio.')
bullet(doc, 'Open Graph (og:) — Meta tags para que cuando compartes el link en redes sociales, salga una vista previa bonita con imagen.')
bullet(doc, 'Twitter Card — Lo mismo pero específico para la red social X (antes Twitter).')
bullet(doc, 'Favicon — El pequeño icono que aparece en la pestaña del navegador.')
bullet(doc, 'Google Fonts — Descarga las fuentes Playfair Display y Lato de internet.')
bullet(doc, 'style.css — Carga los estilos visuales del proyecto.')
bullet(doc, 'Supabase CDN — Carga la librería JavaScript de Supabase.')
bullet(doc, 'supabase.js — Carga el archivo de configuración con la URL y la clave pública de la base de datos.')

h2(doc, '2.2 Header — Barra de navegación fija')

body(doc,
    'El header es la barra oscura que siempre está visible en la parte superior de la página, '
    'sin importar qué tan abajo haya desplazado el usuario. Contiene el logo, el menú de '
    'navegación y el botón de WhatsApp.')

code_block(doc,
'<header id="header">\n'
'  <div class="header-inner">\n'
'    <div class="logo">               ← Logo: imagen + texto\n'
'      <img src="./Icono.jpg" ...>\n'
'      <div class="logo-text">\n'
'        <span class="logo-name">Als Dress</span>\n'
'        <span class="logo-sub">Reynosa · Renta &amp; Venta</span>\n'
'      </div>\n'
'    </div>\n'
'    <button class="menu-toggle" id="menuToggle">  ← Botón hamburguesa (móvil)\n'
'      <span></span><span></span><span></span>\n'
'    </button>\n'
'    <nav id="mainNav">               ← Menú de navegación\n'
'      <a href="#inicio">Inicio</a>\n'
'      <a href="#catalogo">Catálogo</a>\n'
'      ...\n'
'    </nav>\n'
'    <a href="https://wa.me/528991947566" class="header-wa-btn">  ← Botón WA\n'
'      Escríbenos por WhatsApp\n'
'    </a>\n'
'  </div>\n'
'</header>')

body(doc, 'Clases y IDs importantes:')
bullet(doc, 'id="header" — El JavaScript lo usa para agregar la clase "scrolled" cuando el usuario baja la página.')
bullet(doc, 'class="header-inner" — Contenedor con display:flex que alinea horizontalmente todos los elementos.')
bullet(doc, 'id="menuToggle" — Botón que solo aparece en pantallas pequeñas (móvil). Al hacer clic, abre el menú.')
bullet(doc, 'id="mainNav" — El menú de navegación. En móvil se oculta y al hacer clic en el botón hamburguesa, aparece como panel lateral.')
bullet(doc, 'class="header-wa-btn" — Botón con borde blanco transparente en el header.')

h2(doc, '2.3 Hero — Sección principal de portada')

body(doc,
    'El hero es la primera sección grande que ve la clienta cuando entra al sitio. Ocupa toda '
    'la pantalla. Tiene dos partes: el texto a la izquierda y la imagen del vestido a la derecha.')

code_block(doc,
'<section id="inicio" class="hero">\n'
'  <div class="hero-overlay"></div>    ← Capa decorativa (actualmente hidden)\n'
'  <div class="hero-container">        ← Grid de 2 columnas\n'
'\n'
'    <div class="hero-left">           ← COLUMNA IZQUIERDA: texto\n'
'      <span class="hero-pretitle">Als Boutique</span>\n'
'      <p class="hero-eyebrow">Renta de Vestidos en Reynosa</p>\n'
'      <h1 class="hero-title">Brilla en tu próximo <em>evento</em></h1>\n'
'      <div class="hero-deco"></div>   ← Línea decorativa horizontal\n'
'      <p class="hero-description">Vestidos seleccionados...</p>\n'
'      <div class="hero-btns">         ← Botones CTA\n'
'        <a href="#catalogo" class="hero-btn hero-btn-primary">Ver colección</a>\n'
'        <a href="https://wa.me/..." class="hero-btn hero-btn-wa">WhatsApp</a>\n'
'      </div>\n'
'      <div class="hero-trust">        ← 4 puntos de confianza\n'
'        ...\n'
'      </div>\n'
'    </div>\n'
'\n'
'    <div class="hero-right">          ← COLUMNA DERECHA: imagen\n'
'      <div class="hero-image-wrapper">\n'
'        <img src="./hero2.png" alt="Als Dress Hero">\n'
'      </div>\n'
'    </div>\n'
'\n'
'  </div>\n'
'</section>')

body(doc, 'Elementos del hero y su función:')
bullet(doc, 'hero-pretitle — Pequeño texto en mayúsculas arriba del título: "ALS BOUTIQUE".')
bullet(doc, 'hero-eyebrow — Texto descriptivo: "Renta de Vestidos en Reynosa".')
bullet(doc, 'hero-title — El título principal grande. Usa <em> para poner "evento" en cursiva con peso ligero.')
bullet(doc, 'hero-deco — Una línea horizontal delgada, puramente decorativa.')
bullet(doc, 'hero-description — El párrafo de descripción.')
bullet(doc, 'hero-btns — Contenedor con los dos botones: "Ver colección" (oscuro) y "WhatsApp" (verde).')
bullet(doc, 'hero-trust — Cuatro mini-tarjetas con íconos SVG: Renta Fácil, Calidad Garantizada, Luces Increíble, Reynosa.')
bullet(doc, 'hero-right — Columna derecha con la imagen principal. Tiene pseudo-elementos (::before y ::after) en CSS para crear gradientes que fusionan la imagen con el fondo.')
bullet(doc, 'reveal — Clase en varios elementos para que aparezcan con animación al cargar la página.')

h2(doc, '2.4 Sección de Categorías Visuales')

body(doc,
    'Esta sección muestra 4 tarjetas con imágenes: Vestidos, Zapatos, Bolsas y Accesorios. '
    'Cuando la clienta hace clic en una, la página se desplaza al catálogo y activa '
    'automáticamente el filtro de esa categoría.')

code_block(doc,
'<section class="cat-visual-section">\n'
'  <div class="cat-visual-inner">\n'
'    <div class="section-header reveal"> ← Encabezado de sección\n'
'      <p class="section-eyebrow">Explora nuestra colección</p>\n'
'      <h2>Encuentra tu estilo</h2>\n'
'    </div>\n'
'    <div class="cat-cards-grid">         ← Grid de 4 columnas\n'
'\n'
'      <a class="cat-card reveal" data-goto-categoria="Vestido">\n'
'        <img src="./cat-vestidos.png" class="cat-card-img">\n'
'        <div class="cat-card-overlay"></div>   ← Degradado oscuro sobre la imagen\n'
'        <div class="cat-card-content">         ← Texto sobre la imagen\n'
'          <div class="cat-card-icon">...</div> ← Ícono circular con glassmorphism\n'
'          <h3 class="cat-card-name">Vestidos</h3>\n'
'          <span class="cat-card-link">Ver colección →</span>\n'
'        </div>\n'
'      </a>\n'
'\n'
'      ← Se repite para Zapatos, Bolsas, Accesorios...\n'
'\n'
'    </div>\n'
'  </div>\n'
'</section>')

body(doc, 'Atributo especial:')
bullet(doc, 'data-goto-categoria="Vestido" — Atributo personalizado (data attribute). El JavaScript lo lee para saber a qué categoría debe activar al hacer clic.')

h2(doc, '2.5 Sección Catálogo')

body(doc,
    'Esta es la sección más dinámica de la página. Contiene los filtros y el grid de productos '
    'que se llenan automáticamente con datos de Supabase.')

code_block(doc,
'<section id="catalogo" class="catalogo-section">\n'
'\n'
'  <div class="section-header reveal"> ← Encabezado\n'
'    <p class="section-eyebrow">Colección disponible</p>\n'
'    <h2>Catálogo de Temporada</h2>\n'
'  </div>\n'
'\n'
'  <!-- Botones de categoría (ocultos, usados por JS) -->\n'
'  <div class="categorias" id="categorias" style="display:none">\n'
'    <button class="cat-btn active" data-categoria="Vestido">Vestidos</button>\n'
'    <button class="cat-btn" data-categoria="Zapato">Zapatos</button>\n'
'    ...\n'
'  </div>\n'
'\n'
'  <!-- Filtros de disponibilidad y talla -->\n'
'  <div class="filtros reveal" id="filtros">\n'
'    <button class="filtro-btn active" data-filter="todos">Todos</button>\n'
'    <button class="filtro-btn" data-filter="disponible">✓ Disponibles</button>\n'
'    <div class="filtro-separador"></div>\n'
'    <div class="tallas-filtro" id="tallasFiltro"></div>  ← JS inserta botones aquí\n'
'  </div>\n'
'\n'
'  <!-- Grid de productos -->\n'
'  <div id="productos-container" class="productos-grid">\n'
'    <!-- Skeleton loaders: 6 tarjetas grises mientras carga -->\n'
'    <div class="skeleton-card"></div>\n'
'    <div class="skeleton-card"></div>\n'
'    ...\n'
'  </div>\n'
'\n'
'  <p id="mensaje-vacio" class="mensaje-vacio hidden">No se encontraron...</p>\n'
'</section>')

body(doc, 'Elementos clave:')
bullet(doc, 'id="categorias" style="display:none" — Las categorías están ocultas en HTML. Las tarjetas visuales de la sección anterior son las que activan estos botones internamente via JavaScript.')
bullet(doc, 'id="tallasFiltro" — Contenedor vacío que el JavaScript llena con botones de talla dinámicamente, según qué tallas existan en el inventario actual.')
bullet(doc, 'id="productos-container" — El grid donde JavaScript inserta las cards de productos. Inicialmente tiene 6 "skeleton cards" (tarjetas grises animadas) que se muestran mientras cargan los datos reales.')
bullet(doc, 'id="mensaje-vacio" — Mensaje que aparece cuando ningún producto pasa los filtros.')

h2(doc, '2.6 Sección Proceso — "Cómo Funciona"')

body(doc,
    'Esta sección tiene fondo oscuro y explica en 3 pasos cómo funciona la renta: elegir, disfrutar y regresar.')

code_block(doc,
'<section id="proceso" class="proceso-section">\n'
'  <div class="proceso-inner">\n'
'    <div class="section-header reveal">...</div>\n'
'\n'
'    <div class="pasos">             ← Flexbox horizontal con los 3 pasos\n'
'\n'
'      <div class="paso reveal">    ← PASO 1\n'
'        <div class="paso-num-bg">01</div>  ← Número grande decorativo (fondo)\n'
'        <div class="paso-icon-wrap">       ← Ícono circular\n'
'          <svg ...>...</svg>\n'
'        </div>\n'
'        <h3>Elige y Pruébate</h3>\n'
'        <p>Ven a nuestro showroom...</p>\n'
'      </div>\n'
'\n'
'      <div class="paso-conector"></div>    ← Línea conectora con flecha\n'
'\n'
'      <div class="paso reveal">           ← PASO 2: igual estructura\n'
'        ...\n'
'      </div>\n'
'\n'
'      <div class="paso-conector"></div>\n'
'\n'
'      <div class="paso reveal">           ← PASO 3\n'
'        ...\n'
'      </div>\n'
'    </div>\n'
'\n'
'    <div class="proceso-cta reveal">      ← Botón de WhatsApp al final\n'
'      <a href="https://wa.me/..." class="proceso-btn-wa">Agendar mi cita</a>\n'
'    </div>\n'
'  </div>\n'
'</section>')

body(doc, 'Detalle visual:')
bullet(doc, 'paso-num-bg — El número "01", "02", "03" se pone en una fuente muy grande y con muy poca opacidad (5%) en la esquina del paso, creando un efecto de fondo decorativo.')
bullet(doc, 'paso-conector — Una línea horizontal con una pequeña flecha en la punta derecha, creada con CSS usando ::after.')
bullet(doc, 'proceso-section::before — Pseudo-elemento CSS que pone el texto "ALS" gigante y casi invisible en el fondo de toda la sección.')

h2(doc, '2.7 Galería de Clientas')

body(doc,
    'Esta sección muestra fotos de clientas en una cinta que se mueve automáticamente de derecha '
    'a izquierda de forma infinita. La sección está oculta por defecto (clase "hidden") y solo '
    'se hace visible si Supabase devuelve fotos con activa=true.')

code_block(doc,
'<section id="sec-galeria-clientas" class="galeria-clientas-section hidden">\n'
'  <div class="section-header reveal">\n'
'    <p class="section-eyebrow">Ellas ya lo lucieron</p>\n'
'    <h2>Nuestras Clientas</h2>\n'
'  </div>\n'
'  <div class="galeria-cinta-wrapper">   ← Contenedor con overflow:hidden\n'
'    <div id="galeria-clientas-grid" class="galeria-cinta">\n'
'      ← JavaScript inserta las fotos aquí\n'
'    </div>\n'
'  </div>\n'
'</section>')

body(doc, 'Cómo funciona la animación:')
bullet(doc, 'galeria-cinta-wrapper — Oculta el contenido que se sale del contenedor (overflow:hidden) y aplica una máscara de gradiente en los bordes para que las fotos aparezcan y desaparezcan suavemente.')
bullet(doc, 'galeria-cinta — Tiene la animación CSS "cintaScroll" que lo mueve continuamente de 0% a -50%.')
bullet(doc, 'El JavaScript duplica las fotos (pone la lista dos veces) para que cuando la primera copia llega al final, la segunda ya está empezando, creando un loop perfecto.')

h2(doc, '2.8 Sección Testimonios')

body(doc,
    'Muestra opiniones de clientas reales con un resumen de calificación (4.9 estrellas) y '
    'tres tarjetas de reseñas.')

code_block(doc,
'<section id="testimonios" class="testimonios-section">\n'
'  <!-- Barra de calificación -->\n'
'  <div class="testimonios-rating reveal">\n'
'    <div class="rating-score">4.9</div>\n'
'    <div class="rating-mid">\n'
'      <div class="rating-stars">★★★★★</div>\n'
'      <span class="rating-label">Basado en reseñas reales</span>\n'
'    </div>\n'
'    <div class="rating-divider"></div>\n'
'    <div class="rating-badge">Reseñas verificadas</div>\n'
'  </div>\n'
'\n'
'  <!-- Grid de tarjetas -->\n'
'  <div class="testimonios-grid">\n'
'    <div class="testimonio-card reveal">\n'
'      <span class="testimonio-quote-deco">"</span>  ← Comilla decorativa gigante\n'
'      <div class="estrellas">★★★★★</div>\n'
'      <p>Excelente atención y servicio...</p>\n'
'      <div class="testimonio-author">\n'
'        <div class="testimonio-avatar">HC</div>    ← Iniciales del nombre\n'
'        <div class="testimonio-author-info">\n'
'          <h4>Hellen Cruz</h4>\n'
'          <span class="testimonio-verified">Clienta verificada</span>\n'
'        </div>\n'
'      </div>\n'
'    </div>\n'
'    ← Se repiten más tarjetas...\n'
'  </div>\n'
'</section>')

body(doc, 'Para agregar o cambiar un testimonio, simplemente edita el HTML:')
bullet(doc, 'Cambia el texto dentro de <p> para cambiar el mensaje.')
bullet(doc, 'Cambia las letras del testimonio-avatar (ej: "ML" son las iniciales de Mayra López).')
bullet(doc, 'Cambia el nombre dentro de <h4>.')
bullet(doc, 'Para agregar uno nuevo, copia todo el bloque <div class="testimonio-card"> y pégalo.')

h2(doc, '2.9 Sección Ubicación')

body(doc,
    'Muestra dos columnas: a la izquierda una galería de fotos del local con controles de '
    'navegación (flechas y puntos), y a la derecha un mapa de Google Maps embebido.')

code_block(doc,
'<section id="ubicacion" class="ubicacion-section">\n'
'  <div class="ubicacion-grid">       ← Grid de 2 columnas\n'
'\n'
'    <div class="ubicacion-galeria reveal">  ← COLUMNA IZQUIERDA\n'
'      <div class="galeria-wrapper">\n'
'        <img src="./interior1.jpeg" class="foto-slide active">\n'
'        <img src="./interior2.jpeg" class="foto-slide">\n'
'        <img src="./interior3.jpeg" class="foto-slide">\n'
'        <button id="btn-prev" class="ctrl-prev">&#10094;</button>  ← Flecha izq\n'
'        <button id="btn-next" class="ctrl-next">&#10095;</button>  ← Flecha der\n'
'        <div class="galeria-dots" id="galeriaDots"></div>  ← Puntos JS\n'
'      </div>\n'
'    </div>\n'
'\n'
'    <div class="ubicacion-mapa reveal">   ← COLUMNA DERECHA\n'
'      <iframe src="https://www.google.com/maps/embed?..."></iframe>\n'
'    </div>\n'
'\n'
'  </div>\n'
'</section>')

body(doc, 'Cómo funciona el slider de fotos:')
bullet(doc, 'Solo la imagen con clase "foto-slide active" es visible (opacity: 1).')
bullet(doc, 'Al presionar las flechas, JavaScript quita "active" de la foto actual y se la pone a la siguiente.')
bullet(doc, 'Los puntos (dots) también se generan dinámicamente por JavaScript.')
bullet(doc, 'Hay un setInterval que automáticamente avanza cada 6.5 segundos.')

h2(doc, '2.10 Footer y Botones Flotantes')

body(doc,
    'El footer es la parte inferior de la página. Tiene el nombre del negocio, un eslogan y '
    'links a redes sociales (Facebook, Instagram, WhatsApp).')

code_block(doc,
'<footer>\n'
'  <div class="footer-inner">           ← Flexbox: marca a la izquierda, links a la derecha\n'
'    <div class="footer-brand">\n'
'      <h3>Als Dress</h3>\n'
'      <p>Luce increíble en tu próximo evento...</p>\n'
'    </div>\n'
'    <div class="footer-links">\n'
'      <a href="https://web.facebook.com/alsdress...">Facebook</a>\n'
'      <a href="https://www.instagram.com/als_dress">Instagram</a>\n'
'      <a href="https://wa.me/528991947566">WhatsApp</a>\n'
'    </div>\n'
'  </div>\n'
'  <div class="footer-bottom">\n'
'    <p>&copy; 2026 Als Dress · Reynosa, Tamaulipas · ...</p>\n'
'  </div>\n'
'</footer>\n'
'\n'
'<!-- Botones flotantes (siempre visibles en la esquina inferior derecha) -->\n'
'<div class="botones-flotantes">\n'
'  <a href="https://wa.me/..." class="btn-flotante whatsapp">...</a>\n'
'  <a href="https://facebook.com/..." class="btn-flotante facebook">...</a>\n'
'</div>')

body(doc,
    'Los botones flotantes tienen position:fixed, lo que significa que permanecen en la '
    'misma posición en pantalla aunque el usuario desplace la página. Tienen un z-index alto '
    '(950) para aparecer encima de todo el contenido.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CAPÍTULO 3 — CSS
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, 'Capítulo 3: Estilos CSS')

body(doc,
    'CSS (Cascading Style Sheets, u Hojas de Estilos en Cascada) es el lenguaje que controla '
    'cómo se ve la página: colores, tamaños, posiciones, animaciones. El archivo style.css '
    'del proyecto está muy bien organizado en secciones con comentarios. A continuación se '
    'explican los bloques más importantes.')

h2(doc, '3.1 Variables de color (Custom Properties)')

body(doc,
    'Al inicio de style.css, dentro de :root {}, se definen las variables de diseño. '
    ':root es como una "caja de herramientas" global: si defines un color aquí, '
    'puedes usarlo en cualquier parte del CSS escribiendo var(--nombre-variable). '
    'Para cambiar un color en toda la página, solo cambia el valor aquí.')

code_block(doc,
':root {\n'
'  /* Paleta principal — 5 tokens */\n'
'  --rose:       #c0275a;   /* Rosa/fucsia oscuro → color de acento */\n'
'  --rose-pale:  #f2e8ec;   /* Rosa muy claro, casi blanco */\n'
'  --dark:       #1c1c1e;   /* Negro suave (casi negro, tono Apple) */\n'
'  --cream:      #faf8f5;   /* Blanco cálido, fondo principal */\n'
'  --gray:       #8a8a8e;   /* Gris medio, textos secundarios */\n'
'\n'
'  /* Derivados utilitarios */\n'
'  --gray-light: rgba(28,28,30,0.09);  /* Gris muy suave para bordes */\n'
'\n'
'  /* Tipografías */\n'
'  --font-display: \'Playfair Display\', Georgia, serif;\n'
'  --font-body:    \'Lato\', sans-serif;\n'
'\n'
'  /* Espaciado y efectos */\n'
'  --radius:    16px;                               /* Borde redondeado estándar */\n'
'  --shadow:    0 8px 32px rgba(0,0,0,0.07);        /* Sombra suave */\n'
'  --shadow-lg: 0 20px 60px rgba(0,0,0,0.1);        /* Sombra grande */\n'
'  --transition: 0.35s cubic-bezier(0.25,0.46,0.45,0.94);  /* Animación fluida */\n'
'}')

body(doc, 'Cómo usar las variables:')
bullet(doc, 'Para usar el color --rose en un elemento: color: var(--rose);')
bullet(doc, 'Para usar el color de fondo: background: var(--cream);')
bullet(doc, 'Para un borde redondeado: border-radius: var(--radius);')

body(doc, 'Cómo cambiar los colores del sitio completo:')
bullet(doc, 'Cambiar el color de acento (botones, títulos): modifica --rose.')
bullet(doc, 'Cambiar el color de fondo: modifica --cream.')
bullet(doc, 'Cambiar el color oscuro (header, footer, textos principales): modifica --dark.')
bullet(doc, 'Cambiar el color de los textos secundarios: modifica --gray.')

note_box(doc,
    'Los colores se escriben en formato hexadecimal (#c0275a) o RGBA rgba(0,0,0,0.5). '
    'Puedes usar herramientas como Google Color Picker o coolors.co para elegir colores.')

h2(doc, '3.2 Tipografías')

body(doc,
    'El sitio usa dos familias de fuentes con propósitos diferentes:')

bullet(doc, 'Playfair Display (var(--font-display)) — Fuente serif elegante para títulos, nombres de vestidos y texto decorativo. Tiene variante en cursiva que se usa para el "evento" del hero.')
bullet(doc, 'Lato (var(--font-body)) — Fuente sans-serif limpia para el cuerpo del texto, párrafos, botones y etiquetas. Es más fácil de leer en textos largos.')

code_block(doc,
'/* Uso de tipografías */\n'
'.hero-title {\n'
'  font-family: var(--font-display);  /* Playfair Display */\n'
'  font-size: clamp(3.2rem, 5.5vw, 5.8rem);\n'
'  /* clamp(mínimo, ideal, máximo): el tamaño es 5.5% del ancho de pantalla,\n'
'     pero nunca menor a 3.2rem ni mayor a 5.8rem */\n'
'}\n'
'\n'
'body {\n'
'  font-family: var(--font-body);  /* Lato */\n'
'  font-size: 16px;\n'
'  line-height: 1.6;  /* Espaciado entre líneas: 160% del tamaño de fuente */\n'
'}')

body(doc,
    'La función clamp() es muy útil para tipografía responsiva: el tamaño cambia fluidamente '
    'según el ancho de la pantalla, sin necesidad de media queries.')

h2(doc, '3.3 Hero Layout — CSS Grid')

body(doc,
    'El hero usa CSS Grid para dividir la pantalla en dos columnas: texto a la izquierda e '
    'imagen a la derecha.')

code_block(doc,
'.hero-container {\n'
'  display: grid;\n'
'  grid-template-columns: 1fr 52%;\n'
'  /* 1fr = fracción disponible restante (~48%) | 52% = columna de imagen */\n'
'  align-items: stretch;  /* Ambas columnas tienen la misma altura */\n'
'}\n'
'\n'
'.hero-right {\n'
'  position: relative;\n'
'  overflow: hidden;\n'
'  min-height: 520px;\n'
'}\n'
'\n'
'/* Gradiente que fusiona la imagen con el fondo crema */\n'
'.hero-right::before {\n'
'  content: \'\';\n'
'  position: absolute;\n'
'  left: 0; top: 0; bottom: 0;\n'
'  width: 160px;\n'
'  background: linear-gradient(to right, var(--cream) 0%, transparent 100%);\n'
'  z-index: 2;  /* Sobre la imagen, bajo el contenido */\n'
'}')

body(doc,
    'Los pseudo-elementos ::before y ::after son capas virtuales que CSS crea automáticamente. '
    'Se usan para agregar efectos decorativos sin necesidad de elementos HTML adicionales. '
    'El "::before" del hero-right crea el degradado que hace que la imagen se "disuelva" '
    'suavemente en el color de fondo crema.')

h2(doc, '3.4 Cards del Catálogo')

body(doc,
    'Cada tarjeta de producto (.card-producto) tiene un diseño en capas:')

code_block(doc,
'.card-producto {\n'
'  position: relative;\n'
'  background: white;\n'
'  border-radius: 24px;\n'
'  overflow: hidden;  /* Oculta lo que se salga del borde redondeado */\n'
'  display: flex;\n'
'  flex-direction: column;  /* Imagen arriba, info abajo */\n'
'  border: 1px solid var(--gray-light);\n'
'  box-shadow: 0 4px 20px rgba(0,0,0,0.05);\n'
'  opacity: 0;\n'
'  transform: translateY(20px);\n'
'  animation: cardIn 0.5s ease forwards;  /* Animación de entrada */\n'
'}\n'
'\n'
'/* Al pasar el cursor encima */\n'
'.card-producto:hover {\n'
'  transform: translateY(-8px);  /* Se levanta 8px */\n'
'  box-shadow: 0 20px 50px rgba(28,28,30,0.13);  /* Sombra más pronunciada */\n'
'}\n'
'\n'
'/* Animación de entrada cuando aparece la card */\n'
'@keyframes cardIn {\n'
'  from { opacity: 0; transform: translateY(20px); }\n'
'  to   { opacity: 1; transform: translateY(0); }\n'
'}')

body(doc, 'Estructura de capas dentro de la card:')
bullet(doc, '.card-img-wrapper — Contenedor de la imagen con aspect-ratio: 3/4 (proporciones de retrato).')
bullet(doc, '.card-img-wrapper::after — Gradiente oscuro suave en la parte inferior de la imagen.')
bullet(doc, '.badge-estado — Etiqueta en la esquina superior derecha: "Disponible" o "Rentado".')
bullet(doc, '.card-img-hover — Capa invisible que al hacer hover muestra el texto "Ver detalles".')
bullet(doc, '.card-fav-btn — Botón de corazón en la esquina inferior derecha.')
bullet(doc, '.card-info — Sección blanca debajo con nombre, tallas y botón de acción.')

h2(doc, '3.5 Header fijo y menú responsive')

code_block(doc,
'#header {\n'
'  position: fixed;  /* Se queda fijo al hacer scroll */\n'
'  top: 0; left: 0; right: 0;\n'
'  z-index: 1000;    /* Aparece encima de todo */\n'
'  background: var(--dark);\n'
'  transition: box-shadow 0.3s ease;\n'
'}\n'
'\n'
'/* Cuando el usuario hace scroll hacia abajo, JS agrega esta clase */\n'
'#header.scrolled {\n'
'  box-shadow: 0 4px 24px rgba(0,0,0,0.25);\n'
'}\n'
'\n'
'/* En móvil (pantallas menores a 768px) */\n'
'@media (max-width: 768px) {\n'
'  .menu-toggle { display: flex; }  /* Muestra el botón hamburguesa */\n'
'  nav {\n'
'    display: none;\n'
'    position: fixed;\n'
'    top: 0; right: 0;\n'
'    width: min(300px, 80vw);\n'
'    height: 100svh;  /* 100% de la altura visible del viewport */\n'
'    background: var(--dark);\n'
'    flex-direction: column;\n'
'    padding: 80px 28px 28px;\n'
'    box-shadow: -4px 0 30px rgba(0,0,0,0.3);\n'
'  }\n'
'  nav.open { display: flex; }  /* JS agrega "open" al hacer clic */\n'
'}')

h2(doc, '3.6 Botones')

body(doc, 'El sitio tiene varios tipos de botones con estilos distintos:')

code_block(doc,
'/* Botón principal oscuro */\n'
'.hero-btn-primary {\n'
'  background: var(--dark);  /* Fondo oscuro */\n'
'  color: white;\n'
'  box-shadow: 0 4px 20px rgba(28,28,30,0.25);\n'
'}\n'
'\n'
'/* Botón WhatsApp verde */\n'
'.hero-btn-wa {\n'
'  background: #25D366;  /* Verde de WhatsApp */\n'
'  color: white;\n'
'}\n'
'.hero-btn-wa:hover {\n'
'  background: #1aad55;  /* Verde más oscuro al hacer hover */\n'
'}\n'
'\n'
'/* Botón de card disponible */\n'
'.btn-disponible {\n'
'  background: var(--dark);\n'
'  color: white;\n'
'}\n'
'\n'
'/* Botón de card agotado (no hace nada) */\n'
'.btn-agotado {\n'
'  background: var(--cream);\n'
'  color: var(--gray);\n'
'  cursor: not-allowed;  /* Cursor de "no permitido" */\n'
'}')

body(doc,
    'Todos los botones tienen el mismo patrón de hover: transform: translateY(-2px) '
    'que los hace "levitar" ligeramente al pasar el mouse encima, dando una sensación '
    'de profundidad.')

h2(doc, '3.7 Animaciones y transiciones')

body(doc, 'El sitio usa tres tipos de animaciones:')

bullet(doc, '1. Reveal on scroll — Elementos con clase .reveal comienzan invisibles (opacity:0, translateY:30px) y se hacen visibles cuando llegan al viewport.')
bullet(doc, '2. Shimmer skeleton — Las tarjetas de carga tienen un efecto de brillo que se mueve de derecha a izquierda, creando la ilusión de que la página está cargando contenido.')
bullet(doc, '3. Cinta infinita — La galería de clientas se mueve continuamente con una animación CSS llamada cintaScroll.')

code_block(doc,
'/* Reveal on scroll */\n'
'.reveal {\n'
'  opacity: 0;\n'
'  transform: translateY(30px);\n'
'  transition: opacity 0.7s ease, transform 0.7s ease;\n'
'}\n'
'.reveal.visible {\n'
'  opacity: 1;\n'
'  transform: translateY(0);\n'
'}\n'
'\n'
'/* Shimmer skeleton loader */\n'
'@keyframes shimmer {\n'
'  0%   { background-position: 200% 0; }\n'
'  100% { background-position: -200% 0; }\n'
'}\n'
'.skeleton-card {\n'
'  animation: shimmer 1.5s infinite;\n'
'  background: linear-gradient(90deg, var(--cream) 25%, var(--gray-light) 50%, var(--cream) 75%);\n'
'  background-size: 200% 100%;  /* El gradiente es el doble de ancho para moverse */\n'
'}\n'
'\n'
'/* Cinta infinita galería */\n'
'@keyframes cintaScroll {\n'
'  from { transform: translateX(0); }\n'
'  to   { transform: translateX(-50%); }  /* Mueve la mitad (duplicado) */\n'
'}')

h2(doc, '3.8 Responsive — Media Queries')

body(doc,
    'Los media queries son reglas CSS que se aplican solo cuando la pantalla tiene un '
    'determinado tamaño. El diseño se adapta a diferentes dispositivos:')

code_block(doc,
'/* Tablets y pantallas medianas */\n'
'@media (max-width: 980px) {\n'
'  .hero-container {\n'
'    grid-template-columns: 1fr;  /* Una sola columna */\n'
'  }\n'
'  .hero-right {\n'
'    order: 1;  /* La imagen va ARRIBA */\n'
'  }\n'
'  .hero-left {\n'
'    order: 2;  /* El texto va ABAJO */\n'
'    text-align: center;\n'
'  }\n'
'}\n'
'\n'
'/* Teléfonos medianos */\n'
'@media (max-width: 768px) {\n'
'  .menu-toggle { display: flex; }   /* Muestra el botón hamburguesa */\n'
'  nav { display: none; }            /* Oculta el menú horizontal */\n'
'  .productos-grid {\n'
'    grid-template-columns: repeat(2, 1fr);  /* 2 columnas en vez de 3-4 */\n'
'  }\n'
'}\n'
'\n'
'/* Teléfonos pequeños */\n'
'@media (max-width: 480px) {\n'
'  .hero-title { font-size: 2.5rem; }  /* Título más pequeño */\n'
'}')

body(doc, 'Breakpoints usados en el proyecto:')
bullet(doc, '1100px — Oculta el botón de WhatsApp del header.')
bullet(doc, '980px — El hero pasa de 2 columnas a 1 columna (imagen arriba, texto abajo).')
bullet(doc, '900px — Las tarjetas de categorías pasan de 4 a 2 columnas. Los pasos del proceso se apilan verticalmente.')
bullet(doc, '768px — Activa el menú hamburguesa. Las cards del catálogo pasan a 2 columnas.')
bullet(doc, '640px — El lightbox se reorganiza verticalmente (imagen arriba, info abajo).')
bullet(doc, '480px — Ajustes finales para teléfonos muy pequeños.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CAPÍTULO 4 — JAVASCRIPT
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, 'Capítulo 4: Lógica JavaScript — Funciones y Eventos')

body(doc,
    'JavaScript (JS) es el lenguaje de programación del navegador. Se ejecuta en la computadora '
    'del usuario y permite que la página sea interactiva: cargar datos, filtrar productos, '
    'abrir ventanas emergentes (lightbox), animar el menú, etc. '
    'El archivo main.js está organizado en 9 secciones numeradas con comentarios.')

h2(doc, '4.1 Variables globales y constantes')

body(doc,
    'Al inicio del archivo se declaran las constantes y variables que se usan en todo el código:')

code_block(doc,
'// Constantes que no cambian\n'
'const WA_NUMBER  = \'528991947566\';    // Número de WhatsApp (con código de país)\n'
'const APP_NAME   = \'RentaVestidosAPP-250346467\'; // Nombre de la app en AppSheet\n'
'const TABLE_NAME = \'Inventario\';       // Nombre de la tabla en AppSheet\n'
'\n'
'// Estado global — variables que cambian cuando el usuario interactúa\n'
'let todosLosModelos = [];      // Array con todos los productos de Supabase\n'
'let filtroDisponible = false;  // true = mostrar solo disponibles\n'
'let tallaActiva      = null;   // Talla seleccionada (ej: "M") o null si ninguna\n'
'let categoriaActiva  = \'Vestido\'; // Categoría actual')

body(doc, 'Explicación de conceptos:')
bullet(doc, 'const — Declara una constante: un valor que no puede cambiar después de ser declarado.')
bullet(doc, 'let — Declara una variable: un valor que SÍ puede cambiar.')
bullet(doc, '[] — Array vacío: una lista de elementos. todosLosModelos empieza vacío y se llena cuando Supabase responde.')
bullet(doc, 'null — Significa "ningún valor" o "no hay nada seleccionado".')
bullet(doc, 'WA_NUMBER — El número de WhatsApp. Para cambiarlo, edita solo esta línea.')

h2(doc, '4.2 obtenerUrlFoto() — Helper de imágenes')

code_block(doc,
'function obtenerUrlFoto(foto) {\n'
'    const placeholder = \'https://placehold.co/400x500/f5f1eb/8a8a8e?text=Sin+Foto\';\n'
'    if (!foto) return placeholder;          // Si no hay foto, devuelve imagen de reemplazo\n'
'    if (foto.startsWith(\'http\')) return foto; // Si ya es una URL completa, la devuelve tal cual\n'
'    // Si es un nombre de archivo de AppSheet, construye la URL\n'
'    return `https://www.appsheet.com/template/gettablefileurl\n'
'             ?appName=${encodeURIComponent(APP_NAME)}\n'
'             &tableName=${encodeURIComponent(TABLE_NAME)}\n'
'             &fileName=${encodeURIComponent(foto)}`;\n'
'}')

body(doc,
    'Esta función acepta el campo "foto" de un registro de Supabase y lo convierte en una URL '
    'válida para mostrar en la etiqueta <img>. Maneja tres casos: que no haya foto, que ya sea '
    'una URL completa, o que sea un nombre de archivo de AppSheet.')

h2(doc, '4.3 cargarInventario() — Consulta a Supabase')

body(doc,
    'Esta es la función más importante para el catálogo. Se conecta a Supabase y trae '
    'todos los productos publicados.')

code_block(doc,
'async function cargarInventario() {\n'
'    try {\n'
'        // Consulta: todos los registros de "inventario" donde publicado = true,\n'
'        //           ordenados por nombre alfabéticamente\n'
'        const { data, error } = await sb\n'
'            .from(\'inventario\')\n'
'            .select(\'*\')\n'
'            .eq(\'publicado\', true)\n'
'            .order(\'nombre\');\n'
'\n'
'        if (error) throw error;  // Si hay error, lo lanza\n'
'\n'
'        // AGRUPAR: Supabase puede devolver múltiples filas del mismo vestido\n'
'        // (una por cada talla). Esta lógica agrupa todas las tallas bajo un mismo nombre.\n'
'        const modelos = {};\n'
'        data.forEach(item => {\n'
'            const nombre = (item.nombre || \'\').trim();\n'
'            if (!nombre) return;\n'
'\n'
'            if (!modelos[nombre]) {\n'
'                // Primera vez que vemos este nombre: crear objeto base\n'
'                modelos[nombre] = {\n'
'                    ...item,  // Copiar todos los campos\n'
'                    tallasDisponibles: [item.talla].filter(Boolean),\n'
'                    hayDisponible: item.estado_actual === \'Disponible\'\n'
'                };\n'
'            } else {\n'
'                // Ya existe: solo agregar la talla si es nueva\n'
'                if (item.talla && !modelos[nombre].tallasDisponibles.includes(item.talla)) {\n'
'                    modelos[nombre].tallasDisponibles.push(item.talla);\n'
'                }\n'
'                // Si alguna talla está disponible, marcar el modelo como disponible\n'
'                if (item.estado_actual === \'Disponible\') modelos[nombre].hayDisponible = true;\n'
'            }\n'
'        });\n'
'\n'
'        todosLosModelos = Object.values(modelos);  // Convertir a array\n'
'        construirFiltrosTallas();  // Crear los botones de talla\n'
'        renderizarCatalogo();      // Mostrar los productos\n'
'\n'
'    } catch (err) {\n'
'        console.error(\'Error cargando catálogo:\', err);\n'
'        // Mostrar mensaje de error en lugar del catálogo\n'
'        document.getElementById(\'productos-container\').innerHTML =\n'
'            \'<p class="mensaje-vacio">No se pudo cargar el catálogo.</p>\';\n'
'    }\n'
'}')

body(doc, 'Conceptos clave usados en esta función:')
bullet(doc, 'async/await — Palabras clave para manejar operaciones que toman tiempo (como llamadas a internet). "await" pausa la función hasta que Supabase responde, sin bloquear el navegador.')
bullet(doc, 'try/catch — Bloque para manejar errores. Si algo falla dentro del "try", el "catch" lo captura y muestra un mensaje al usuario.')
bullet(doc, 'spread operator (...item) — Copia todos los campos del objeto "item" en el nuevo objeto.')
bullet(doc, 'Object.values(modelos) — Convierte el objeto agrupado en un array (lista).')

h2(doc, '4.4 renderizarCatalogo() — Mostrar productos')

body(doc,
    'Esta función genera el HTML de cada card de producto y lo inserta en el contenedor del catálogo. '
    'Se llama cada vez que cambia un filtro.')

code_block(doc,
'function renderizarCatalogo() {\n'
'    const contenedor = document.getElementById(\'productos-container\');\n'
'    const msgVacio   = document.getElementById(\'mensaje-vacio\');\n'
'\n'
'    let lista = [...todosLosModelos];  // Copia del array global\n'
'\n'
'    // APLICAR FILTROS en orden:\n'
'    // 1. Filtro por categoría\n'
'    if (categoriaActiva) {\n'
'        lista = lista.filter(m => (m.tipo || \'Vestido\') === categoriaActiva);\n'
'    }\n'
'    // 2. Filtro por disponibilidad\n'
'    if (filtroDisponible) lista = lista.filter(m => m.hayDisponible);\n'
'    // 3. Filtro por talla\n'
'    if (tallaActiva) lista = lista.filter(m => m.tallasDisponibles.includes(tallaActiva));\n'
'\n'
'    contenedor.innerHTML = \'\';  // Limpiar el grid actual\n'
'\n'
'    if (lista.length === 0) {\n'
'        // Sin resultados: verificar si la categoría tiene productos\n'
'        const tieneProductos = todosLosModelos.some(m => (m.tipo||\'Vestido\') === categoriaActiva);\n'
'        if (!tieneProductos) {\n'
'            // Mostrar tarjeta "Próximamente"\n'
'            contenedor.innerHTML = `<div class="proximamente-card">...</div>`;\n'
'        } else {\n'
'            msgVacio.classList.remove(\'hidden\');  // Mostrar "No se encontraron"\n'
'        }\n'
'        return;\n'
'    }\n'
'\n'
'    // Crear una card por cada producto filtrado\n'
'    lista.forEach((vestido, i) => {\n'
'        const urlFoto  = obtenerUrlFoto(vestido.foto);\n'
'        const disponible = vestido.hayDisponible;\n'
'        const card = document.createElement(\'div\');  // Crear elemento div\n'
'        card.className = \'card-producto\';\n'
'        card.style.animationDelay = `${i * 60}ms`;  // Delay progresivo\n'
'        card.innerHTML = `...`;  // HTML de la card\n'
'        // Agregar event listeners\n'
'        card.querySelector(\'.card-img-wrapper\').addEventListener(\'click\', () =>\n'
'            abrirLightbox(vestido, urlFoto));\n'
'        contenedor.appendChild(card);  // Insertar en el DOM\n'
'    });\n'
'}')

body(doc, 'El método .filter() crea un nuevo array con solo los elementos que cumplen la condición:')
code_block(doc,
'// Ejemplo: filtrar solo los vestidos disponibles\n'
'lista = lista.filter(m => m.hayDisponible);\n'
'// m => m.hayDisponible es una "arrow function":\n'
'// para cada modelo "m", devuelve si hayDisponible es true')

h2(doc, '4.5 crearLightbox() y abrirLightbox()')

body(doc,
    'El lightbox es el modal (ventana emergente) que aparece cuando la clienta hace clic en '
    'una card del catálogo. Muestra la foto en grande y los detalles del vestido.')

body(doc, 'crearLightbox() — Se ejecuta UNA sola vez al inicio y construye el HTML del modal:')

code_block(doc,
'function crearLightbox() {\n'
'    const lb = document.createElement(\'div\');\n'
'    lb.id = \'lightbox\';\n'
'    lb.innerHTML = `\n'
'        <div class="lightbox-inner">\n'
'            <div class="lightbox-img-wrap">\n'
'                <img id="lightbox-img" src="" alt="">\n'
'            </div>\n'
'            <div class="lightbox-info">\n'
'                <button class="lightbox-cerrar" id="lightboxCerrar">✕</button>\n'
'                <p class="lb-eyebrow"><span id="lightbox-categoria"></span></p>\n'
'                <h2 class="lightbox-nombre" id="lightbox-nombre"></h2>\n'
'                ...\n'
'            </div>\n'
'        </div>`;\n'
'    document.body.appendChild(lb);  // Agrega el modal al final del body\n'
'    // Cerrar con el botón ✕\n'
'    document.getElementById(\'lightboxCerrar\').addEventListener(\'click\', cerrarLightbox);\n'
'    // Cerrar al hacer clic fuera del modal\n'
'    lb.addEventListener(\'click\', (e) => { if (e.target === lb) cerrarLightbox(); });\n'
'    // Cerrar al presionar Escape\n'
'    document.addEventListener(\'keydown\', (e) => { if (e.key === \'Escape\') cerrarLightbox(); });\n'
'}')

body(doc, 'abrirLightbox() — Se ejecuta cada vez que se hace clic en una card, y llena el modal con los datos del vestido seleccionado:')

code_block(doc,
'function abrirLightbox(vestido, urlFoto) {\n'
'    const lb = document.getElementById(\'lightbox\');\n'
'    const disponible = vestido.hayDisponible;\n'
'\n'
'    // Llenar los datos del vestido en el modal\n'
'    document.getElementById(\'lightbox-img\').src = urlFoto;\n'
'    document.getElementById(\'lightbox-nombre\').textContent = vestido.nombre;\n'
'    document.getElementById(\'lightbox-categoria\').textContent =\n'
'        (vestido.tipo || \'Vestido\').toUpperCase();\n'
'\n'
'    // Crear chips de tallas\n'
'    document.getElementById(\'lightbox-tallas\').innerHTML =\n'
'        vestido.tallasDisponibles.sort()\n'
'        .map(t => `<span class="talla-chip">${t}</span>`).join(\'\');\n'
'\n'
'    // Configurar el botón según disponibilidad\n'
'    const btnWa = document.getElementById(\'lightbox-btn-wa\');\n'
'    if (disponible) {\n'
'        const textoWA = `Hola Als Dress! Vi el vestido "${vestido.nombre}". ¿Hay disponibilidad?`;\n'
'        btnWa.onclick = () => window.open(`https://wa.me/${WA_NUMBER}?text=${encodeURIComponent(textoWA)}`, \'_blank\');\n'
'        btnWa.innerHTML = `${WA_ICON_SVG} Consultar por WhatsApp`;\n'
'    } else {\n'
'        btnWa.innerHTML = \'No disponible por el momento\';\n'
'        btnWa.onclick = null;\n'
'    }\n'
'\n'
'    lb.classList.add(\'open\');         // Muestra el lightbox\n'
'    document.body.style.overflow = \'hidden\'; // Impide el scroll de fondo\n'
'}')

h2(doc, '4.6 cargarGaleriaClientas()')

body(doc,
    'Carga las fotos de clientas desde Supabase y crea la cinta animada:')

code_block(doc,
'async function cargarGaleriaClientas() {\n'
'    const contenedor = document.getElementById(\'galeria-clientas-grid\');\n'
'\n'
'    // Consulta: fotos activas, ordenadas de más nueva a más antigua\n'
'    const { data, error } = await sb\n'
'        .from(\'galeria_clientas\')\n'
'        .select(\'*\')\n'
'        .eq(\'activa\', true)\n'
'        .order(\'created_at\', { ascending: false });\n'
'\n'
'    if (error || !data || data.length === 0) {\n'
'        // Si no hay fotos, ocultar la sección completa\n'
'        document.getElementById(\'sec-galeria-clientas\').classList.add(\'hidden\');\n'
'        return;\n'
'    }\n'
'\n'
'    document.getElementById(\'sec-galeria-clientas\').classList.remove(\'hidden\');\n'
'\n'
'    // Función auxiliar para crear cada elemento de la cinta\n'
'    const crearItem = (foto) => {\n'
'        const div = document.createElement(\'div\');\n'
'        div.className = \'galeria-cliente-item\';\n'
'        div.innerHTML = `<img src="${foto.foto_url}" loading="lazy" ...>`;\n'
'        div.onclick = () => abrirLightboxClientas(foto.foto_url, foto.nombre);\n'
'        return div;\n'
'    };\n'
'\n'
'    // TRUCO: Duplicar las fotos para el loop infinito\n'
'    [...data, ...data].forEach(foto => contenedor.appendChild(crearItem(foto)));\n'
'\n'
'    // Ajustar velocidad: más fotos = animación más lenta\n'
'    contenedor.style.animationDuration = `${Math.max(30, data.length * 5)}s`;\n'
'}')

body(doc,
    'El truco del loop infinito: la animación CSS mueve la cinta desde 0% hasta -50% (la mitad). '
    'Si el contenedor tiene 20 fotos (10 originales + 10 duplicadas), cuando llega al final '
    'de las primeras 10, visualmente ya estamos viendo las 10 originales de nuevo. '
    'La animación entonces "reinicia" sin que se note el salto.')

h2(doc, '4.7 Filtros de categoría y talla')

code_block(doc,
'// Filtro por categoría (Vestidos, Zapatos, etc.)\n'
'document.querySelectorAll(\'[data-categoria]\').forEach(btn => {\n'
'    btn.addEventListener(\'click\', () => {\n'
'        // Quitar "active" de todos los botones\n'
'        document.querySelectorAll(\'[data-categoria]\').forEach(b => b.classList.remove(\'active\'));\n'
'        // Poner "active" en el presionado\n'
'        btn.classList.add(\'active\');\n'
'        categoriaActiva = btn.dataset.categoria;  // Actualizar estado global\n'
'        tallaActiva = null;  // Resetear filtro de talla\n'
'        construirFiltrosTallas();  // Reconstruir botones de talla\n'
'        renderizarCatalogo();      // Actualizar grid\n'
'    });\n'
'});\n'
'\n'
'// construirFiltrosTallas() — Crea dinámicamente los botones de talla\n'
'function construirFiltrosTallas() {\n'
'    // Filtrar modelos de la categoría activa\n'
'    const modelosFiltrados = todosLosModelos.filter(m =>\n'
'        !categoriaActiva || (m.tipo || \'Vestido\') === categoriaActiva\n'
'    );\n'
'    // Recolectar todas las tallas únicas en un Set\n'
'    const tallaSet = new Set();\n'
'    modelosFiltrados.forEach(m => m.tallasDisponibles.forEach(t => tallaSet.add(t)));\n'
'\n'
'    // Ordenar: primero tallas con nombre (XXS, XS, S...) luego numéricas\n'
'    const orden = [\'XXS\',\'2XS\',\'XS\',\'S\',\'M\',\'L\',\'XL\',\'2XL\',\'3XL\',\'4XL\'];\n'
'    const tallas = [...tallaSet].sort((a, b) => {\n'
'        const ia = orden.indexOf(a), ib = orden.indexOf(b);\n'
'        if (ia !== -1 && ib !== -1) return ia - ib;\n'
'        ...\n'
'    });\n'
'\n'
'    // Crear un botón por cada talla\n'
'    const contenedor = document.getElementById(\'tallasFiltro\');\n'
'    contenedor.innerHTML = \'\';\n'
'    tallas.forEach(t => {\n'
'        const btn = document.createElement(\'button\');\n'
'        btn.className = \'filtro-btn\';\n'
'        btn.textContent = t;\n'
'        btn.dataset.talla = t;\n'
'        btn.addEventListener(\'click\', () => toggleTalla(t, btn));\n'
'        contenedor.appendChild(btn);\n'
'    });\n'
'}')

h2(doc, '4.8 Event listeners del menú y scroll')

code_block(doc,
'// Referencias a elementos del DOM\n'
'const header     = document.getElementById(\'header\');\n'
'const menuToggle = document.getElementById(\'menuToggle\');\n'
'const mainNav    = document.getElementById(\'mainNav\');\n'
'\n'
'// Agregar clase "scrolled" cuando el usuario baja 20px\n'
'window.addEventListener(\'scroll\', () => {\n'
'    header.classList.toggle(\'scrolled\', window.scrollY > 20);\n'
'}, { passive: true }); // passive:true mejora el rendimiento del scroll\n'
'\n'
'// Abrir/cerrar menú hamburguesa\n'
'menuToggle.addEventListener(\'click\', () => {\n'
'    const isOpen = mainNav.classList.toggle(\'open\');  // toggle devuelve true/false\n'
'    menuToggle.classList.toggle(\'open\', isOpen);\n'
'    document.body.style.overflow = isOpen ? \'hidden\' : \'\';\n'
'    // Si está abierto: bloquea el scroll del fondo\n'
'    // Si está cerrado: restaura el scroll\n'
'});\n'
'\n'
'// Cerrar el menú cuando se hace clic en un link\n'
'mainNav.querySelectorAll(\'a\').forEach(link => {\n'
'    link.addEventListener(\'click\', () => {\n'
'        mainNav.classList.remove(\'open\');\n'
'        menuToggle.classList.remove(\'open\');\n'
'        document.body.style.overflow = \'\';\n'
'    });\n'
'});')

h2(doc, '4.9 Reveal on scroll — IntersectionObserver')

code_block(doc,
'// IntersectionObserver: observa si un elemento entra en el viewport\n'
'const revealObserver = new IntersectionObserver((entries) => {\n'
'    entries.forEach((entry, i) => {\n'
'        if (entry.isIntersecting) {  // Si el elemento es visible\n'
'            setTimeout(() => {\n'
'                entry.target.classList.add(\'visible\');  // Activar animación\n'
'            }, i * 80);  // Delay escalonado: 0ms, 80ms, 160ms...\n'
'            revealObserver.unobserve(entry.target);  // Dejar de observar\n'
'        }\n'
'    });\n'
'}, { threshold: 0.12 });\n'
'// threshold:0.12 = activa cuando el 12% del elemento es visible\n'
'\n'
'// Observar todos los elementos con clase "reveal"\n'
'document.querySelectorAll(\'.reveal\').forEach(el => revealObserver.observe(el));')

body(doc,
    'IntersectionObserver es una API del navegador que detecta cuándo un elemento '
    'entra o sale del área visible de la pantalla (el "viewport"). Es más eficiente '
    'que los viejos métodos basados en el evento scroll.')

h2(doc, '4.10 DOMContentLoaded — Inicialización')

code_block(doc,
'document.addEventListener(\'DOMContentLoaded\', () => {\n'
'    // Este código se ejecuta cuando el HTML está completamente cargado\n'
'\n'
'    // 1. Actualizar todos los links de WhatsApp con el número correcto\n'
'    const waUrl = `https://wa.me/${WA_NUMBER}`;\n'
'    document.querySelectorAll(`a[href*="wa.me"]`).forEach(a => a.href = waUrl);\n'
'\n'
'    // 2. Ocultar botones flotantes cuando el hero es visible\n'
'    const botonesFlotantes = document.querySelector(\'.botones-flotantes\');\n'
'    const heroEl = document.getElementById(\'inicio\');\n'
'    if (botonesFlotantes && heroEl) {\n'
'        new IntersectionObserver(([entry]) => {\n'
'            botonesFlotantes.classList.toggle(\'oculto\', entry.isIntersecting);\n'
'        }, { threshold: 0.2 }).observe(heroEl);\n'
'    }\n'
'\n'
'    // 3. Tarjetas de categorías → scroll + activar filtro\n'
'    document.querySelectorAll(\'[data-goto-categoria]\').forEach(card => {\n'
'        card.addEventListener(\'click\', (e) => {\n'
'            e.preventDefault();\n'
'            const cat = card.dataset.gotoCategoria;\n'
'            document.getElementById(\'catalogo\').scrollIntoView({ behavior: \'smooth\' });\n'
'            setTimeout(() => {\n'
'                const btn = document.querySelector(`[data-categoria="${cat}"]`);\n'
'                if (btn) btn.click();\n'
'            }, 600);  // Espera 600ms a que termine el scroll\n'
'        });\n'
'    });\n'
'\n'
'    // 4. Crear el lightbox (solo la estructura HTML, sin datos)\n'
'    crearLightbox();\n'
'\n'
'    // 5. Cargar productos del catálogo desde Supabase\n'
'    cargarInventario();\n'
'\n'
'    // 6. Cargar galería de clientas desde Supabase\n'
'    cargarGaleriaClientas();\n'
'});')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CAPÍTULO 5 — GUÍA DE CAMBIOS COMUNES
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, 'Capítulo 5: Guía de Cambios Comunes')

body(doc,
    'Esta sección es práctica y directa. Para cada tarea común, se indica exactamente '
    'qué archivo abrir y qué línea cambiar.')

separator(doc)

h2(doc, '5.1 Cambiar el número de WhatsApp')

body(doc,
    'El número de WhatsApp se define en una sola constante al inicio de main.js. '
    'Cambiándola ahí, se actualiza automáticamente en todos los botones de la página.')

body(doc, 'Archivo: main.js — Línea 5')
code_block(doc,
'// ANTES:\n'
'const WA_NUMBER = \'528991947566\';\n'
'\n'
'// DESPUÉS (ejemplo con un número nuevo):\n'
'const WA_NUMBER = \'528991234567\';\n'
'\n'
'// El formato es: código de país (52 para México) + número sin guiones ni espacios')

note_box(doc,
    'El script en DOMContentLoaded también actualiza todos los <a href="wa.me/..."> del HTML '
    'con el número definido en WA_NUMBER. Sin embargo, por buena práctica, también actualiza '
    'los links directamente en index.html buscando "528991947566" y reemplazándolo.')

separator(doc)
h2(doc, '5.2 Cambiar el color de acento (rosa)')

body(doc, 'Archivo: style.css — Línea 9')
code_block(doc,
'/* ANTES: */\n'
'--rose: #c0275a;  /* Rosa oscuro */\n'
'\n'
'/* DESPUÉS: Cambiar a otro color, por ejemplo morado: */\n'
'--rose: #7c3aed;\n'
'\n'
'/* O azul marino: */\n'
'--rose: #1e40af;')

body(doc,
    'Esta variable controla principalmente los títulos h2 en el documento Word generado. '
    'En la página web, --rose se usa en algunos detalles. El color principal de la marca '
    'son los botones oscuros (var(--dark)), así que para cambiar el color de los botones '
    'hay que modificar --dark también.')

body(doc, 'Para cambiar el color principal de los botones y el header:')
code_block(doc,
'/* ANTES: */\n'
'--dark: #1c1c1e;\n'
'\n'
'/* DESPUÉS: Cambiar a color corporativo, ej. burdeos: */\n'
'--dark: #6b0f2a;')

separator(doc)
h2(doc, '5.3 Cambiar textos del Hero')

body(doc, 'Archivo: index.html — sección <section id="inicio" class="hero">')
code_block(doc,
'<!-- Texto pequeño arriba del título -->\n'
'<span class="hero-pretitle">Als Boutique</span>\n'
'↓ Cambia "Als Boutique" por lo que desees\n'
'\n'
'<!-- Subtítulo -->\n'
'<p class="hero-eyebrow">Renta de Vestidos en Reynosa</p>\n'
'↓ Cambia la ciudad o el texto descriptivo\n'
'\n'
'<!-- Título principal (H1) -->\n'
'<h1 class="hero-title">\n'
'  Brilla en tu<br>\n'
'  próximo <em>evento</em>\n'
'</h1>\n'
'↓ <em> pone en cursiva la palabra "evento"\n'
'   <br> fuerza un salto de línea\n'
'\n'
'<!-- Descripción -->\n'
'<p class="hero-description">\n'
'  Vestidos seleccionados para graduaciones, galas y eventos especiales.\n'
'  Luce increíble sin gastar una fortuna.\n'
'</p>')

separator(doc)
h2(doc, '5.4 Cambiar la imagen del Hero')

body(doc,
    'La imagen principal del hero es el archivo hero2.png en la raíz del proyecto.')

bullet(doc, 'Paso 1: Prepara la nueva imagen. Recomendamos PNG o JPEG de al menos 1200px de alto, con la persona centrada y fondo liso o con poco detalle en los bordes.')
bullet(doc, 'Paso 2: Guarda la imagen con el mismo nombre "hero2.png" en la carpeta del proyecto, reemplazando la anterior.')
bullet(doc, 'Paso 3: Si quieres usar otro nombre, busca en index.html: <img src="./hero2.png"> y cambia "hero2.png" por el nombre de tu archivo.')

code_block(doc,
'<!-- ANTES: -->\n'
'<img src="./hero2.png" alt="Als Dress Hero">\n'
'\n'
'<!-- DESPUÉS (imagen con otro nombre): -->\n'
'<img src="./mi-nueva-imagen.jpg" alt="Als Dress Hero">')

separator(doc)
h2(doc, '5.5 Agregar un vestido al catálogo')

body(doc,
    'Los vestidos se gestionan desde Supabase, NO desde el código HTML. Para agregar uno:')

bullet(doc, 'Paso 1: Abre tu navegador y ve a supabase.com.')
bullet(doc, 'Paso 2: Inicia sesión y abre el proyecto de Als Dress.')
bullet(doc, 'Paso 3: En el menú de la izquierda, haz clic en "Table Editor".')
bullet(doc, 'Paso 4: Selecciona la tabla "inventario".')
bullet(doc, 'Paso 5: Haz clic en "Insert row" (Insertar fila).')
bullet(doc, 'Paso 6: Llena los campos: nombre (ej: "Vestido Azul Medianoche"), tipo ("Vestido"), talla ("M"), estado_actual ("Disponible"), foto (URL de la imagen o nombre del archivo), publicado (true).')
bullet(doc, 'Paso 7: Guarda. El vestido aparecerá automáticamente en el catálogo la próxima vez que se recargue la página.')

note_box(doc,
    'Si el mismo vestido existe en varias tallas, crea UNA FILA POR TALLA, con el mismo nombre. '
    'El JavaScript agrupa automáticamente todas las tallas del mismo nombre en una sola card.')

separator(doc)
h2(doc, '5.6 Cambiar las imágenes de las tarjetas de categorías')

body(doc,
    'Las imágenes de las 4 tarjetas (Vestidos, Zapatos, Bolsas, Accesorios) son archivos locales.')

code_block(doc,
'<!-- En index.html, sección .cat-visual-section: -->\n'
'<img src="./cat-vestidos.png" alt="Vestidos" class="cat-card-img">\n'
'<img src="./cat-zapatos.png" alt="Zapatos" class="cat-card-img">\n'
'<img src="./cat-bolsas.png" alt="Bolsas" class="cat-card-img">\n'
'<img src="./cat-accesorios.png" alt="Accesorios" class="cat-card-img">')

bullet(doc, 'Reemplaza el archivo de imagen correspondiente manteniendo el mismo nombre, O')
bullet(doc, 'Cambia el atributo src en el HTML por la ruta de tu nueva imagen.')
bullet(doc, 'Recomendamos imágenes cuadradas o en proporción 3:4 (retrato), al menos 600x800px.')

separator(doc)
h2(doc, '5.7 Cambiar las fotos del local (sección Ubicación)')

code_block(doc,
'<!-- En index.html, sección #ubicacion: -->\n'
'<img src="./interior1.jpeg" class="foto-slide active" alt="Interior Als Dress 1">\n'
'<img src="./interior2.jpeg" class="foto-slide" alt="Interior Als Dress 2">\n'
'<img src="./interior3.jpeg" class="foto-slide" alt="Interior Als Dress 3">')

bullet(doc, 'Para reemplazar una foto: guarda la nueva imagen con el mismo nombre (interior1.jpeg, etc.) y reemplaza el archivo.')
bullet(doc, 'Para agregar una foto nueva: copia una línea <img> y cambia el src y el alt. El slider de JavaScript funciona con cualquier cantidad de imágenes.')
bullet(doc, 'Para quitar una foto: simplemente elimina la línea <img> correspondiente.')

separator(doc)
h2(doc, '5.8 Cambiar los testimonios')

body(doc, 'Los testimonios están directamente en el HTML. Para cambiar uno, edita:')

code_block(doc,
'<div class="testimonio-card reveal">\n'
'  <span class="testimonio-quote-deco">"</span>\n'
'  <div class="estrellas">★★★★★</div>\n'
'\n'
'  <!-- TEXTO DEL TESTIMONIO: -->\n'
'  <p>Excelente atención y servicio, 100% recomendado, tienen mucha variedad de vestidos.</p>\n'
'\n'
'  <div class="testimonio-author">\n'
'    <!-- INICIALES DEL AVATAR: -->\n'
'    <div class="testimonio-avatar">HC</div>\n'
'    <div class="testimonio-author-info">\n'
'      <!-- NOMBRE: -->\n'
'      <h4>Hellen Cruz</h4>\n'
'      <span class="testimonio-verified">Clienta verificada</span>\n'
'    </div>\n'
'  </div>\n'
'</div>')

separator(doc)
h2(doc, '5.9 Cambiar el mapa de Google Maps')

body(doc,
    'El mapa está incrustado con un <iframe>. Para cambiarlo a otra ubicación:')

bullet(doc, 'Paso 1: Ve a Google Maps (maps.google.com) y busca la nueva dirección.')
bullet(doc, 'Paso 2: Haz clic en "Compartir".')
bullet(doc, 'Paso 3: Selecciona la pestaña "Incorporar un mapa".')
bullet(doc, 'Paso 4: Copia el código HTML del iframe.')
bullet(doc, 'Paso 5: En index.html, reemplaza el <iframe> existente en la sección #ubicacion.')

separator(doc)
h2(doc, '5.10 Agregar una sección nueva a la página')

body(doc,
    'Para agregar una nueva sección (por ejemplo, una sección de preguntas frecuentes):')

bullet(doc, 'Paso 1: Abre index.html.')
bullet(doc, 'Paso 2: Encuentra dónde quieres insertar la sección (entre los comentarios <!-- ===== NOMBRE ===== -->).')
bullet(doc, 'Paso 3: Copia la estructura base de una sección existente y modifica su contenido.')
bullet(doc, 'Paso 4: Asigna un id único a la sección (ej: id="faq").')
bullet(doc, 'Paso 5: Agrega el link al menú en el header si quieres que sea navegable.')
bullet(doc, 'Paso 6: Agrega los estilos necesarios en style.css.')

code_block(doc,
'<!-- Ejemplo de nueva sección básica: -->\n'
'<section id="faq" class="faq-section">\n'
'  <div class="section-header reveal">\n'
'    <p class="section-eyebrow">Tus preguntas</p>\n'
'    <h2>Preguntas Frecuentes</h2>\n'
'  </div>\n'
'  <div class="faq-contenido">\n'
'    <!-- Tu contenido aquí -->\n'
'  </div>\n'
'</section>\n'
'\n'
'<!-- Agregar al menú en header: -->\n'
'<a href="#faq">FAQ</a>\n'
'\n'
'/* En style.css: */\n'
'.faq-section {\n'
'  padding: 100px 24px;\n'
'  max-width: 1200px;\n'
'  margin: 0 auto;\n'
'}')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CAPÍTULO 6 — GLOSARIO
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, 'Capítulo 6: Glosario de Términos Técnicos')

body(doc,
    'Este glosario explica los términos técnicos más frecuentes en el código del proyecto, '
    'en lenguaje claro y con ejemplos del propio sitio.')

glosario_items = [
    ('API', 'Application Programming Interface (Interfaz de Programación de Aplicaciones). '
     'Es un "contrato" que define cómo dos sistemas de software se comunican entre sí. '
     'En el proyecto, usamos la API de Supabase para consultar la base de datos, '
     'y la API de WhatsApp (wa.me) para abrir conversaciones.'),

    ('Array', 'Una lista ordenada de elementos en JavaScript. '
     'En el proyecto, "todosLosModelos" es un array que guarda todos los vestidos cargados. '
     'Se escribe entre corchetes: [elemento1, elemento2, ...]. '
     'Se puede recorrer con .forEach(), filtrar con .filter(), etc.'),

    ('async/await', 'Sintaxis de JavaScript para manejar operaciones asíncronas (que toman tiempo). '
     '"async" declara que una función puede tener operaciones lentas. '
     '"await" pausa la ejecución hasta que esa operación termina, sin bloquear el navegador. '
     'Se usa en cargarInventario() y cargarGaleriaClientas() para esperar la respuesta de Supabase.'),

    ('Breakpoint', 'Punto de quiebre en CSS responsivo. Es el ancho de pantalla en el que '
     'el diseño cambia para adaptarse al dispositivo. En el proyecto los breakpoints son '
     '1100px, 980px, 900px, 768px, 640px y 480px.'),

    ('CDN', 'Content Delivery Network (Red de Distribución de Contenido). '
     'Una red de servidores distribuidos geográficamente que entrega archivos de forma más rápida. '
     'Se usa para cargar la librería de Supabase y las fuentes de Google Fonts desde un CDN, '
     'en vez de guardar esos archivos en el proyecto.'),

    ('CSS', 'Cascading Style Sheets. Lenguaje que controla la presentación visual de la página web: '
     'colores, tipografías, tamaños, posiciones, animaciones. '
     'El archivo style.css contiene todos los estilos del proyecto.'),

    ('DOM', 'Document Object Model. La representación de la página HTML como un árbol de objetos '
     'que JavaScript puede leer y modificar. Cuando JavaScript hace '
     'document.getElementById("catalogo"), está accediendo al DOM para encontrar un elemento.'),

    ('Event Listener', 'Función que "escucha" un evento (clic, tecla presionada, scroll) y '
     'ejecuta código cuando ese evento ocurre. En el proyecto, los botones del menú, '
     'las cards del catálogo y los filtros tienen event listeners.'),

    ('Flexbox', 'Módulo de CSS para alinear y distribuir elementos en una sola dimensión '
     '(fila o columna). El header usa display:flex para alinear el logo, el menú y el '
     'botón de WhatsApp horizontalmente.'),

    ('Grid', 'CSS Grid Layout. Sistema de diseño bidimensional (filas Y columnas). '
     'El hero usa grid de 2 columnas, el catálogo usa grid de auto-fill, '
     'y la sección de testimonios usa grid de minmax.'),

    ('HTML', 'HyperText Markup Language. Lenguaje de marcado que define la estructura y '
     'el contenido de la página web. Usa etiquetas como <section>, <div>, <h1>, <p>. '
     'El archivo index.html contiene toda la estructura del proyecto.'),

    ('JavaScript', 'Lenguaje de programación del navegador. Permite crear páginas interactivas: '
     'cargar datos dinámicamente, responder a clics, filtrar contenido, hacer animaciones. '
     'El archivo main.js contiene toda la lógica del proyecto.'),

    ('Lightbox', 'Ventana emergente (modal) que aparece sobre el contenido de la página '
     'para mostrar una imagen o información en detalle. En el proyecto, hacer clic en '
     'una card del catálogo abre el lightbox con la foto grande y los detalles del vestido.'),

    ('Media Query', 'Regla CSS que se aplica solo cuando la pantalla cumple cierta condición, '
     'como tener un ancho máximo. Ej: @media (max-width: 768px) { ... } aplica estilos '
     'solo en pantallas de hasta 768px de ancho (teléfonos celulares).'),

    ('Netlify', 'Plataforma de hosting para sitios web estáticos. Despliega automáticamente '
     'el sitio de Als Dress desde el repositorio de código. La URL del sitio es '
     'alsdressrentadevestidos.netlify.app.'),

    ('Pseudo-elemento', 'Elemento virtual que CSS crea automáticamente sin necesitar HTML extra. '
     '::before crea un elemento antes del contenido, ::after después. '
     'Se usan para gradientes decorativos en el hero y las secciones oscuras.'),

    ('Responsive', 'Diseño que se adapta a diferentes tamaños de pantalla '
     '(teléfonos, tablets, computadoras). El sitio de Als Dress es responsivo: '
     'en móvil el hero se apila verticalmente, el menú se convierte en hamburguesa, '
     'y el catálogo muestra 2 columnas en vez de 3 o 4.'),

    ('Supabase', 'Plataforma de base de datos en la nube (Backend as a Service). '
     'Proporciona una base de datos PostgreSQL, una API REST automática y autenticación. '
     'En el proyecto guarda el inventario de vestidos y la galería de clientas.'),

    ('SVG', 'Scalable Vector Graphics. Formato de imagen basado en código que escala perfectamente '
     'a cualquier tamaño sin perder calidad. Todos los íconos del proyecto (WhatsApp, '
     'corazón, calendario, ubicación) son SVGs incrustados directamente en el HTML.'),

    ('Variable CSS', 'Valor reutilizable definido en :root con el prefijo --. '
     'Ej: --rose: #c0275a. Se usa como var(--rose) en cualquier propiedad CSS. '
     'Facilita cambiar el diseño completo modificando un solo lugar.'),

    ('z-index', 'Propiedad CSS que controla qué elemento aparece "encima" cuando se superponen. '
     'A mayor z-index, más al frente. El header tiene z-index:1000, el lightbox z-index:1000, '
     'los botones flotantes z-index:950.'),

    ('IntersectionObserver', 'API del navegador que detecta cuándo un elemento entra o sale '
     'del viewport (área visible). Se usa para las animaciones de reveal: cuando un elemento '
     'con clase .reveal entra al viewport, se le agrega la clase .visible y aparece animado.'),

    ('Template Literal', 'Sintaxis de JavaScript con backticks (`) que permite insertar '
     'variables dentro de strings. Ej: `Hola ${nombre}!` inserta el valor de la variable nombre. '
     'Se usa extensivamente para construir el HTML de las cards y el lightbox.'),

    ('encodeURIComponent', 'Función de JavaScript que convierte caracteres especiales a '
     'formato URL-safe. Se usa para el texto del mensaje de WhatsApp, ya que los espacios, '
     'acentos y signos de interrogación deben codificarse: "¿Tienen disponibilidad?" '
     'se convierte en "%C2%BFTienen%20disponibilidad%3F".'),

    ('classList.toggle', 'Método de JavaScript que agrega una clase si no existe, o la quita si '
     'ya existe. Ej: header.classList.toggle("scrolled", window.scrollY > 20) '
     'agrega la clase "scrolled" cuando scrollY > 20 y la quita cuando no.'),
]

for term, definition in glosario_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r_term = p.add_run(term + ' — ')
    r_term.font.name = 'Calibri'
    r_term.font.size = Pt(11)
    r_term.font.bold = True
    r_term.font.color.rgb = RGBColor(0xc0, 0x27, 0x5a)
    r_def = p.add_run(definition)
    r_def.font.name = 'Calibri'
    r_def.font.size = Pt(11)
    r_def.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  NOTAS FINALES
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, 'Notas Finales y Recursos')

h2(doc, 'Estructura resumida del flujo de la aplicación')

code_block(doc,
'CARGA DE PÁGINA\n'
'│\n'
'├── HTML/CSS se renderiza (header, hero, skeleton cards...)\n'
'│\n'
'├── main.js se ejecuta:\n'
'│   ├── Configura event listeners (scroll, menú, filtros)\n'
'│   └── Configura IntersectionObserver para reveals\n'
'│\n'
'└── DOMContentLoaded dispara:\n'
'    ├── crearLightbox()         → Crea estructura del modal\n'
'    ├── cargarInventario()      → Consulta Supabase\n'
'    │   ├── Agrupa por nombre\n'
'    │   ├── construirFiltrosTallas()\n'
'    │   └── renderizarCatalogo()  → Inserta cards en el DOM\n'
'    └── cargarGaleriaClientas() → Consulta Supabase\n'
'        └── Si hay fotos: muestra sección + cinta animada\n'
'\n'
'INTERACCIÓN DE LA USUARIA\n'
'│\n'
'├── Clic en categoría visual → scroll + btn.click() → renderizarCatalogo()\n'
'├── Clic en filtro talla   → toggleTalla() → renderizarCatalogo()\n'
'├── Clic en card           → abrirLightbox() → llena modal\n'
'├── Clic "Consultar WA"    → window.open(wa.me + texto) → WhatsApp\n'
'└── Clic menú hamburguesa  → classList.toggle("open")')

h2(doc, 'Recursos de aprendizaje recomendados')

body(doc, 'Para seguir aprendiendo sobre las tecnologías usadas en este proyecto:')

bullet(doc, 'HTML y CSS básico: developer.mozilla.org (MDN Web Docs, en español)')
bullet(doc, 'JavaScript desde cero: javascript.info (en inglés, muy completo)')
bullet(doc, 'CSS Flexbox visual: flexboxfroggy.com (juego interactivo en español)')
bullet(doc, 'CSS Grid visual: cssgridgarden.com (juego interactivo)')
bullet(doc, 'Supabase documentación: supabase.com/docs')
bullet(doc, 'Netlify documentación: docs.netlify.com')
bullet(doc, 'Colores y paletas: coolors.co, colorhunt.co')
bullet(doc, 'Íconos SVG gratuitos: heroicons.com, lucide.dev')
bullet(doc, 'Google Fonts: fonts.google.com')

h2(doc, 'Contacto y soporte técnico')

body(doc,
    'Este sitio fue desarrollado para Als Dress, Reynosa, Tamaulipas. '
    'Para soporte técnico o modificaciones avanzadas, contactar al desarrollador '
    'que configuró el proyecto originalmente.')

body(doc,
    'Para cambios simples (textos, colores, imágenes), esta guía contiene todos los '
    'pasos necesarios. Para cambios más complejos (nueva funcionalidad, integración '
    'con otros sistemas), se recomienda contar con un desarrollador web.')

p_final = doc.add_paragraph()
p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_final.paragraph_format.space_before = Pt(40)
r_final = p_final.add_run('— Fin del Manual de Código — Als Dress · Reynosa, Tamaulipas · 2026 —')
r_final.font.name = 'Calibri'
r_final.font.size = Pt(10)
r_final.font.italic = True
r_final.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)

# ── Guardar ──────────────────────────────────────────────────────────────────
output_path = '/home/user/Als/manual-codigo.docx'
doc.save(output_path)
print(f'Documento generado exitosamente: {output_path}')
